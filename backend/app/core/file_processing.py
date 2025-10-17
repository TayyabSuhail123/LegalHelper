"""File processing and text extraction services."""

import os
import tempfile
import magic
import aiofiles
from typing import Optional, Tuple
from pathlib import Path

import PyPDF2
from docx import Document
from fastapi import UploadFile, HTTPException

from app.models.file_processing import FileType, UploadedFile, ProcessingStatus
from app.core.storage import FileStorageManager


class FileProcessingService:
    """Service for handling file upload and text extraction."""
    
    # Allowed MIME types
    ALLOWED_MIME_TYPES = {
        "application/pdf": FileType.PDF,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": FileType.DOCX,
        "text/plain": FileType.TXT,
    }
    
    def __init__(self, upload_dir: str = "./uploads", max_file_size: int = 50 * 1024 * 1024):
        """
        Initialize the file processing service.
        
        Args:
            upload_dir: Directory for temporary file storage
            max_file_size: Maximum allowed file size in bytes
        """
        self.max_file_size = max_file_size
        self.storage_manager = FileStorageManager(upload_dir)
    
    @property
    def MAX_FILE_SIZE(self) -> int:
        """Backward compatibility property."""
        return self.max_file_size
    
    async def validate_file(self, file: UploadFile) -> Tuple[bool, str, Optional[FileType]]:
        """
        Validate uploaded file.
        
        Args:
            file: Uploaded file object
            
        Returns:
            Tuple of (is_valid, error_message, file_type)
        """
        # Check file size
        if file.size and file.size > self.max_file_size:
            return False, f"File size ({file.size} bytes) exceeds maximum allowed size ({self.max_file_size} bytes)", None
        
        # Read file content for MIME type detection
        content = await file.read()
        await file.seek(0)  # Reset file pointer
        
        # Detect MIME type
        mime_type = magic.from_buffer(content, mime=True)
        
        if mime_type not in self.ALLOWED_MIME_TYPES:
            return False, f"Unsupported file type: {mime_type}. Allowed types: {list(self.ALLOWED_MIME_TYPES.keys())}", None
        
        file_type = self.ALLOWED_MIME_TYPES[mime_type]
        
        # Additional filename extension check
        file_ext = Path(file.filename).suffix.lower()
        expected_extensions = {
            FileType.PDF: [".pdf"],
            FileType.DOCX: [".docx"],
            FileType.TXT: [".txt"],
        }
        
        if file_ext not in expected_extensions[file_type]:
            return False, f"File extension {file_ext} doesn't match detected type {file_type}", None
        
        return True, "File validation successful", file_type

    async def extract_text(self, file_path: str, file_type: FileType) -> str:
        """
        Extract text from uploaded file.
        
        Args:
            file_path: Path to the uploaded file
            file_type: Type of the file
            
        Returns:
            Extracted text content
            
        Raises:
            HTTPException: If text extraction fails
        """
        try:
            if file_type == FileType.PDF:
                return await self._extract_pdf_text(file_path)
            elif file_type == FileType.DOCX:
                return await self._extract_docx_text(file_path)
            elif file_type == FileType.TXT:
                return await self._extract_txt_text(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Failed to extract text from {file_type} file: {str(e)}"
            )
    
    async def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file."""
        text = ""
        
        with open(file_path, "rb") as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
        
        if not text.strip():
            raise ValueError("No text content found in PDF file")
        
        return text.strip()
    
    async def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        doc = Document(file_path)
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        if not text.strip():
            raise ValueError("No text content found in DOCX file")
        
        return text.strip()
    
    async def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from TXT file."""
        async with aiofiles.open(file_path, "r", encoding="utf-8") as file:
            text = await file.read()
        
        if not text.strip():
            raise ValueError("No text content found in TXT file")
        
        return text.strip()
    
    async def cleanup_file(self, file_path: str) -> None:
        """Remove uploaded file from disk."""
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
        except Exception as e:
            # Log error but don't raise - file cleanup shouldn't break the flow
            print(f"Warning: Failed to cleanup file {file_path}: {e}")
    
    async def process_file(self, file: UploadFile) -> UploadedFile:
        """
        Complete file processing pipeline.
        
        Args:
            file: Uploaded file object
            
        Returns:
            UploadedFile model with processing results
        """
        # Validate file
        is_valid, error_message, file_type = await self.validate_file(file)
        if not is_valid:
            raise HTTPException(status_code=400, detail=error_message)
        
        # Create file record
        uploaded_file = UploadedFile(
            filename=file.filename,
            file_type=file_type,
            file_size=file.size or 0,
            processing_status=ProcessingStatus.PROCESSING
        )
        
        try:
            # Read file content
            file_content = await file.read()
            
            # Store file temporarily
            file_path = await self.storage_manager.store_temp_file(
                file_content, file.filename, uploaded_file.file_id
            )
            
            # Extract text
            extracted_text = await self.extract_text(str(file_path), file_type)
            
            # Update file record
            uploaded_file.extracted_text = extracted_text
            uploaded_file.processing_status = ProcessingStatus.COMPLETED
            
            # Note: File will be automatically cleaned up by storage manager
            
            return uploaded_file
            
        except Exception as e:
            uploaded_file.processing_status = ProcessingStatus.FAILED
            uploaded_file.error_message = str(e)
            
            # Clean up file on error
            await self.storage_manager.remove_file(uploaded_file.file_id)
            
            raise HTTPException(
                status_code=500,
                detail="File processing failed. Please try again."
            )
