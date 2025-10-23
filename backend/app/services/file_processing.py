"""File processing and text extraction services."""

import logging
import os
import traceback
from pathlib import Path

import aiofiles
import magic
import PyPDF2
from docx import Document
from fastapi import HTTPException, UploadFile

from app.core.config import Settings
from app.models.file_processing import DocumentFile, FileType, ProcessingStatus
from app.services.storage import FileStorageManager

# Configure logging
logger = logging.getLogger(__name__)


class FileProcessingService:
    """Service for handling file upload and text extraction."""

    # Allowed MIME types
    ALLOWED_MIME_TYPES = {
        "application/pdf": FileType.PDF,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": FileType.DOCX,
        "text/plain": FileType.TXT,
    }

    def __init__(
        self,
        upload_dir: str = "./uploads",
        max_file_size: int = 50 * 1024 * 1024,
        settings: Settings | None = None,
    ):
        """
        Initialize the file processing service.

        Args:
            upload_dir: Directory for temporary file storage
            max_file_size: Maximum allowed file size in bytes
            settings: Application settings instance
        """
        from app.core.config import settings as global_settings

        self.settings = settings or global_settings
        self.upload_dir = upload_dir
        self.max_file_size = max_file_size
        self.storage_manager = FileStorageManager(
            base_dir=upload_dir, cleanup_interval=self.settings.file_cleanup_interval
        )

    @property
    def MAX_FILE_SIZE(self) -> int:
        """Backward compatibility property."""
        return self.max_file_size

    async def validate_file(self, file: UploadFile) -> tuple[bool, str, FileType | None]:
        """
        Validate uploaded file.

        Args:
            file: Uploaded file object

        Returns:
            Tuple of (is_valid, error_message, file_type)
        """
        try:
            logger.info(f"Validating file: {file.filename}")
            logger.debug(f"File content type: {file.content_type}")

            # Check file size
            if file.size and file.size > self.max_file_size:
                error_msg = f"File size ({file.size} bytes) exceeds maximum allowed size ({self.max_file_size} bytes)"
                logger.warning(f"File validation failed - {error_msg}")
                return False, error_msg, None

            # Read file content for MIME type detection
            logger.debug("Reading file content for MIME type detection")
            content = await file.read()
            await file.seek(0)  # Reset file pointer

            # Detect MIME type
            logger.debug("Detecting MIME type")
            mime_type = magic.from_buffer(content, mime=True)
            logger.info(f"Detected MIME type: {mime_type}")

            if mime_type not in self.ALLOWED_MIME_TYPES:
                error_msg = f"Unsupported file type: {mime_type}. Allowed types: {list(self.ALLOWED_MIME_TYPES.keys())}"
                logger.warning(f"File validation failed - {error_msg}")
                return False, error_msg, None

            file_type = self.ALLOWED_MIME_TYPES[mime_type]
            logger.info(f"File type determined: {file_type}")

            # Additional filename extension check - be more flexible
            file_ext = Path(file.filename).suffix.lower()
            allowed_extensions = [".pdf", ".docx", ".txt", ".doc", ".md"]

            if file_ext not in allowed_extensions:
                error_msg = f"File extension {file_ext} is not supported. Allowed extensions: {allowed_extensions}"
                logger.warning(f"File validation failed - {error_msg}")
                return False, error_msg, None

            logger.info(f"File validation successful for: {file.filename}")
            return True, "File validation successful", file_type

        except Exception as e:
            logger.error(f"Exception during file validation for file: {file.filename}")
            logger.error(f"Error type: {type(e).__name__}")
            logger.error(f"Error message: {str(e)}")
            logger.error(f"Traceback: {traceback.format_exc()}")
            return False, f"File validation failed: {str(e)}", None

    async def extract_text(self, file_path: str, file_type: FileType) -> str:
        """
        Extract text from uploaded file.

        Args:
            file_path: Path to the uploaded file
            file_type: Type of the file

        Returns:
            Extracted text content

        Raises:
            HTTPException: If text extraction fails
        """
        try:
            logger.info(f"Starting text extraction from {file_type} file: {file_path}")

            if file_type == FileType.PDF:
                text = await self._extract_pdf_text(file_path)
            elif file_type == FileType.DOCX:
                text = await self._extract_docx_text(file_path)
            elif file_type == FileType.TXT:
                text = await self._extract_txt_text(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")

            logger.info(f"Text extraction successful. Extracted {len(text)} characters")
            return text

        except Exception as e:
            logger.error(f"Text extraction failed for {file_type} file: {file_path}")
            logger.error(f"Error type: {type(e).__name__}")
            logger.error(f"Error message: {str(e)}")
            logger.error(f"Traceback: {traceback.format_exc()}")

            # Provide more user-friendly error messages
            if file_type == FileType.PDF and "No text content found" in str(e):
                detail = "The PDF appears to be image-based or scanned. Please try a text-based PDF or a different file format."
            elif file_type == FileType.DOCX and "No text content found" in str(e):
                detail = "The DOCX file appears to be empty or contains only images. Please try a document with text content."
            elif file_type == FileType.TXT and "No text content found" in str(e):
                detail = "The text file appears to be empty. Please try a file with content."
            elif "encoding" in str(e).lower() or "decode" in str(e).lower():
                detail = "The file has encoding issues. Please try saving it in UTF-8 format or use a different file."
            else:
                detail = f"Failed to extract text from {file_type} file. Please try a different file or format."

            raise HTTPException(status_code=500, detail=detail)

    async def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file."""
        try:
            logger.debug(f"Extracting PDF text from: {file_path}")
            text = ""

            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                logger.debug(f"PDF has {total_pages} pages")

                for page_num in range(total_pages):
                    logger.debug(f"Processing page {page_num + 1}/{total_pages}")
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    text += page_text + "\n"

            if not text.strip():
                logger.warning(f"No text content found in PDF file: {file_path}")
                logger.info("This might be a scanned/image-based PDF that requires OCR")
                raise ValueError(
                    "No text content found in PDF file. This might be a scanned document that requires OCR processing."
                )

            logger.debug(f"PDF text extraction completed. Total characters: {len(text)}")
            return text.strip()

        except Exception as e:
            logger.error(f"PDF text extraction failed for file: {file_path}")
            logger.error(f"Error details: {str(e)}")
            raise

    async def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            logger.debug(f"Extracting DOCX text from: {file_path}")
            doc = Document(file_path)
            text = ""

            # Extract text from paragraphs
            paragraph_count = len(doc.paragraphs)
            logger.debug(f"DOCX has {paragraph_count} paragraphs")
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # Also extract text from tables
            table_count = len(doc.tables)
            logger.debug(f"DOCX has {table_count} tables")
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"

            if not text.strip():
                logger.warning(f"No text content found in DOCX file: {file_path}")
                raise ValueError("No text content found in DOCX file")

            logger.debug(f"DOCX text extraction completed. Total characters: {len(text)}")
            return text.strip()

        except Exception as e:
            logger.error(f"DOCX text extraction failed for file: {file_path}")
            logger.error(f"Error details: {str(e)}")
            raise

    async def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from TXT file."""
        try:
            logger.debug(f"Extracting TXT text from: {file_path}")

            # Try multiple encodings
            encodings = ["utf-8", "utf-16", "latin-1", "cp1252"]
            text = None

            for encoding in encodings:
                try:
                    async with aiofiles.open(file_path, encoding=encoding) as file:
                        text = await file.read()
                    logger.debug(f"Successfully read TXT file with {encoding} encoding")
                    break
                except UnicodeDecodeError:
                    logger.debug(f"Failed to read with {encoding} encoding, trying next...")
                    continue

            if text is None:
                raise ValueError(
                    f"Could not decode TXT file with any of the supported encodings: {encodings}"
                )

            if not text.strip():
                logger.warning(f"No text content found in TXT file: {file_path}")
                raise ValueError("No text content found in TXT file")

            logger.debug(f"TXT text extraction completed. Total characters: {len(text)}")
            return text.strip()

        except Exception as e:
            logger.error(f"TXT text extraction failed for file: {file_path}")
            logger.error(f"Error details: {str(e)}")
            raise
            raise ValueError("No text content found in TXT file")

        return text.strip()

    async def cleanup_file(self, file_path: str) -> None:
        """Remove uploaded file from disk."""
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
        except Exception as e:
            # Log error but don't raise - file cleanup shouldn't break the flow
            print(f"Warning: Failed to cleanup file {file_path}: {e}")

    async def process_file(self, file: UploadFile) -> DocumentFile:
        """
        Complete file processing pipeline.

        Args:
            file: Uploaded file object

        Returns:
            UploadedFile model with processing results
        """
        # Validate file
        is_valid, error_message, file_type = await self.validate_file(file)
        if not is_valid:
            raise HTTPException(status_code=400, detail=error_message)

        # Create file record
        uploaded_file = DocumentFile(
            filename=file.filename,
            file_type=file_type,
            file_size=file.size or 0,
            processing_status=ProcessingStatus.PROCESSING,
        )

        try:
            # Read file content
            file_content = await file.read()

            # Store file temporarily
            file_path = await self.storage_manager.store_temp_file(
                file_content, file.filename, uploaded_file.file_id
            )

            # Extract text
            extracted_text = await self.extract_text(str(file_path), file_type)

            # Update file record
            uploaded_file.extracted_text = extracted_text
            uploaded_file.processing_status = ProcessingStatus.COMPLETED

            # Note: File will be automatically cleaned up by storage manager

            return uploaded_file

        except Exception as e:
            uploaded_file.processing_status = ProcessingStatus.FAILED
            uploaded_file.error_message = str(e)

            # Clean up file on error
            await self.storage_manager.remove_file(uploaded_file.file_id)

            raise HTTPException(status_code=500, detail="File processing failed. Please try again.")
